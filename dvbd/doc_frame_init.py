# Document Frame Initializer
from docx import Document
import tkinter as tk
from tkinter.filedialog import askopenfilename
from inputter import file_browse
import os


def new_word(img, filler):
    new_doc = Document()
    new_doc.add_heading("New Word document with image.")
    usedwd = os.getcwd() + "\\" + str(filler)
    img.savefig(usedwd)
    new_doc.add_picture(usedwd + '.png')
    os.remove(usedwd + '.png')
    new_doc.save(usedwd + '.docx')

def doc_frame(fig_prof):
    doc_frame = tk.Tk()
    doc_frame.geometry("500x500")
    doc_frame.title("Uploader Option Window")
    file_select_button = tk.Button(doc_frame, text = "Select File",
                            command = lambda : file_browse("Print"))
    file_select_button.place(x = 50, y = 50)
    new_word_button = tk.Button(doc_frame, text = "New Word File",
                                command = lambda : new_word(fig_prof.initial_img, fig_prof.filler))
    new_word_button.place(x = 110, y = 50)
    
    
